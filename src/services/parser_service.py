# =============================================================================
# FILE: src/services/parser_service.py (UPGRADED VERSION)
# =============================================================================
from pathlib import Path
import structlog
from typing import Any, Dict, List, Optional
import pandas as pd
import fitz  # PyMuPDF (imported as fitz) - Superior to PyPDF2[web:175][web:185]
from bs4 import BeautifulSoup
import json
import docx  # python-docx
import pytesseract  # OCR for images[web:189]
from PIL import Image
import io
import asyncio

logger = structlog.get_logger()

class ParserService:
    """Advanced multi-format file parser with OCR support and memory protection"""
    
    def __init__(self):
        # Configuration for memory protection
        self.max_rows_per_file = 10000  # Maximum rows for CSV/Excel/HTML tables
        self.max_columns_per_file = 100  # Maximum columns for CSV/Excel
        self.max_file_size_mb = 50  # Maximum file size in MB
        self.max_pdf_pages = 100  # Maximum PDF pages to process
    
    async def parse_file(self, filepath: Path) -> Dict[str, Any]:
        """Parse file based on extension with comprehensive format support and safety limits"""
        suffix = filepath.suffix.lower()
        logger.info("parsing_file", filepath=str(filepath), type=suffix)
        
        # Check file size before processing
        file_size_mb = filepath.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            logger.warning("file_size_exceeded", filepath=str(filepath), size_mb=file_size_mb, max_mb=self.max_file_size_mb)
            return {
                "error": f"File size {file_size_mb:.1f}MB exceeds maximum {self.max_file_size_mb}MB",
                "type": suffix.lstrip('.'),
                "dataframe": pd.DataFrame(),  # Always return empty dataframe for consistency
                "data": []
            }
        
        # Mapping of extensions to parser methods
        parsers = {
            '.pdf': self._parse_pdf_advanced,
            '.csv': self._parse_csv,
            '.tsv': self._parse_csv,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.json': self._parse_json,
            '.jsonl': self._parse_jsonl,
            '.html': self._parse_html_advanced,
            '.htm': self._parse_html_advanced,
            '.txt': self._parse_text,
            '.docx': self._parse_docx,
            '.xml': self._parse_xml,
            '.yaml': self._parse_yaml,
            '.yml': self._parse_yaml,
            # Image formats (OCR)
            '.png': self._parse_image_ocr,
            '.jpg': self._parse_image_ocr,
            '.jpeg': self._parse_image_ocr,
            '.tiff': self._parse_image_ocr,
            '.bmp': self._parse_image_ocr
        }
        
        parser = parsers.get(suffix, self._parse_text)
        try:
            result = await parser(filepath)
            # Ensure consistent return format with dataframe when applicable
            return self._normalize_result(result, suffix.lstrip('.'))
        except Exception as e:
            logger.error("parse_file_error", error=str(e), filepath=str(filepath))
            return self._create_error_result(str(e), suffix.lstrip('.'))
    
    def _normalize_result(self, result: Dict[str, Any], file_type: str) -> Dict[str, Any]:
        """Ensure consistent result format with dataframe and data fields"""
        normalized = result.copy()
        normalized.setdefault("type", file_type)
        
        # Always include dataframe field (empty if not applicable)
        if "dataframe" not in normalized:
            if "data" in normalized and isinstance(normalized["data"], list) and normalized["data"]:
                try:
                    normalized["dataframe"] = pd.DataFrame(normalized["data"])
                except Exception as e:
                    logger.warning("dataframe_creation_failed", error=str(e))
                    normalized["dataframe"] = pd.DataFrame()
            else:
                normalized["dataframe"] = pd.DataFrame()
        
        # Ensure data field exists for tabular data
        if "data" not in normalized and "dataframe" in normalized and not normalized["dataframe"].empty:
            try:
                normalized["data"] = normalized["dataframe"].to_dict(orient="records")
            except Exception as e:
                logger.warning("data_creation_failed", error=str(e))
                normalized["data"] = []
        
        return normalized
    
    def _create_error_result(self, error: str, file_type: str) -> Dict[str, Any]:
        """Create consistent error result format"""
        return {
            "error": error,
            "type": file_type,
            "dataframe": pd.DataFrame(),  # Always include empty dataframe
            "data": []  # Always include empty data list
        }
    
    async def _parse_pdf_advanced(self, filepath: Path) -> Dict[str, Any]:
        """
        Parse PDF using PyMuPDF with table extraction and page limits.
        PyMuPDF is superior to PyPDF2 for text and table extraction[web:175][web:185][web:187].
        """
        try:
            doc = fitz.open(filepath)
            full_text = ""
            pages_data = []
            all_tables = []
            
            # Limit number of pages processed
            max_pages = min(len(doc), self.max_pdf_pages)
            if len(doc) > self.max_pdf_pages:
                logger.warning("pdf_page_limit", total_pages=len(doc), max_pages=self.max_pdf_pages)
            
            for page_num, page in enumerate(doc, 1):
                if page_num > max_pages:
                    break
                    
                # Extract text
                page_text = page.get_text()
                full_text += f"\n--- Page {page_num} ---\n{page_text}"
                
                # Extract tables using PyMuPDF's table detection[web:187]
                tables_on_page = page.find_tables()
                page_tables = []
                
                for table in tables_on_page:
                    # Extract table as pandas DataFrame
                    table_data = table.extract()
                    if table_data:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        # Apply row limits to individual tables
                        if len(df) > self.max_rows_per_file:
                            logger.warning("pdf_table_row_limit", table_rows=len(df), max_rows=self.max_rows_per_file)
                            df = df.head(self.max_rows_per_file)
                        
                        page_tables.append({
                            'data': df.to_dict(orient='records'),
                            'dataframe': df,
                            'bbox': table.bbox
                        })
                        all_tables.append(df)
                
                pages_data.append({
                    'page': page_num,
                    'text': page_text,
                    'tables': page_tables,
                    'table_count': len(page_tables)
                })
            
            doc.close()
            
            # Combine all tables into one dataframe if multiple tables exist
            combined_df = pd.DataFrame()
            if all_tables:
                try:
                    combined_df = pd.concat(all_tables, ignore_index=True)
                    # Apply overall row limit to combined dataframe
                    if len(combined_df) > self.max_rows_per_file:
                        logger.warning("pdf_combined_table_limit", combined_rows=len(combined_df), max_rows=self.max_rows_per_file)
                        combined_df = combined_df.head(self.max_rows_per_file)
                except Exception as e:
                    logger.warning("pdf_table_combine_failed", error=str(e))
                    combined_df = all_tables[0] if all_tables else pd.DataFrame()
            
            logger.info(
                "PDF parsed with PyMuPDF",
                pages=len(pages_data),
                total_tables=len(all_tables)
            )
            
            return {
                'type': 'pdf',
                'num_pages': len(pages_data),
                'full_text': full_text,
                'pages': pages_data,
                'all_tables': [df.to_dict(orient='records') for df in all_tables],
                'tables_dataframes': all_tables,
                'dataframe': combined_df,  # Main dataframe for analysis
                'data': combined_df.to_dict(orient='records') if not combined_df.empty else []
            }
            
        except Exception as e:
            logger.error("pdf_parse_error", error=str(e))
            return {'error': str(e), 'type': 'pdf'}
    
    async def _parse_csv(self, filepath: Path) -> Dict[str, Any]:
        """Parse CSV/TSV with better error handling and row/column limits"""
        try:
            # Auto-detect delimiter
            delimiter = '\t' if filepath.suffix == '.tsv' else ','
            
            # Try different encodings with CORRECT pandas parameters
            df = None
            for encoding in ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']:
                try:
                    # FIXED: Use correct parameter names and read full file with limits
                    df = pd.read_csv(
                        filepath, 
                        delimiter=delimiter, 
                        encoding=encoding,
                        # CORRECT PARAMETER NAMES:
                        on_bad_lines='skip',           # Skip problematic lines
                        encoding_errors='replace',     # Replace invalid characters
                        nrows=self.max_rows_per_file,  # Apply row limit upfront
                        dtype=str,                     # Read as strings initially
                        low_memory=False               # Reduce memory warnings
                    )
                    break
                except (UnicodeDecodeError, pd.errors.ParserError) as e:
                    logger.debug(f"CSV parse failed with encoding {encoding}: {e}")
                    continue
            
            # FIXED: Add fallback with python engine for problematic files
            if df is None:
                try:
                    logger.warning("Using python engine fallback for CSV parsing")
                    df = pd.read_csv(
                        filepath,
                        delimiter=delimiter,
                        encoding='utf-8',
                        engine='python',               # More robust but slower engine
                        on_bad_lines='skip',
                        encoding_errors='replace',
                        nrows=self.max_rows_per_file,
                        dtype=str
                    )
                except Exception as fallback_error:
                    logger.error("CSV python engine fallback also failed", error=str(fallback_error))
                    return {
                        'type': 'csv',
                        'shape': (0, 0),
                        'columns': [],
                        'data': [],
                        'summary': {},
                        'dataframe': pd.DataFrame(),
                        'warning': 'CSV file could not be parsed with any method',
                        'error': str(fallback_error)
                    }
            
            # FIXED: Apply column limits directly during read
            if len(df.columns) > self.max_columns_per_file:
                logger.warning("csv_column_limit", columns=len(df.columns), max_columns=self.max_columns_per_file)
                # Keep only first N columns
                df = df.iloc[:, :self.max_columns_per_file]
            
            # FIXED: Convert numeric columns where possible (preserve your original logic)
            for col in df.columns:
                try:
                    # First try direct numeric conversion
                    df[col] = pd.to_numeric(df[col], errors='ignore')
                    
                    # If that fails but column looks numeric, try cleaning
                    if df[col].dtype == 'object':
                        # Remove common non-numeric characters
                        clean_series = df[col].astype(str).str.replace(r'[^\d.-]', '', regex=True)
                        numeric_series = pd.to_numeric(clean_series, errors='ignore')
                        # Only replace if we got meaningful conversion
                        if numeric_series.notna().sum() > len(df) * 0.5:  # >50% converted
                            df[col] = numeric_series
                except Exception as conv_error:
                    logger.debug(f"Column {col} numeric conversion failed: {conv_error}")
                    continue
            
            # FIXED: Handle empty DataFrame case
            if df is None or df.empty:
                return {
                    'type': 'csv',
                    'shape': (0, 0),
                    'columns': [],
                    'data': [],
                    'summary': {},
                    'dataframe': pd.DataFrame(),
                    'warning': 'Empty or unreadable CSV file'
                }
            
            # FIXED: Final row limit check (should already be applied from read_csv nrows)
            if len(df) > self.max_rows_per_file:
                logger.warning("csv_row_limit_truncated", rows=len(df), max_rows=self.max_rows_per_file)
                df = df.head(self.max_rows_per_file)
            
            # FIXED: Create data sample for large files (preserve your structure but add sampling)
            data_records = df.to_dict(orient='records')
            
            # For very large datasets, provide sampled data but keep full dataframe
            if len(data_records) > 1000:
                sampled_data = data_records[:100]  # Sample first 100 records for response
                data_warning = f"Data sampled to 100 records from {len(data_records)} total"
            else:
                sampled_data = data_records
                data_warning = None
            
            result = {
                'type': 'csv',
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'data': sampled_data,
                'summary': df.describe(include='all').to_dict() if not df.empty else {},
                'dataframe': df
            }
            
            # Add warning if data was sampled
            if data_warning:
                result['warning'] = data_warning
                
            return result
            
        except Exception as e:
            logger.error("csv_parse_error", error=str(e), filepath=str(filepath))
            return {
                'error': str(e), 
                'type': 'csv',
                'dataframe': pd.DataFrame(),
                'data': [],
                'shape': (0, 0),
                'columns': []
            }
    async def _parse_excel(self, filepath: Path) -> Dict[str, Any]:
        """Parse Excel with all sheets and row limits"""
        try:
            xl_file = pd.ExcelFile(filepath)
            sheets = {}
            all_dataframes = {}
            primary_df = pd.DataFrame()
            
            for sheet_name in xl_file.sheet_names:
                try:
                    # Read with row limit
                    df = pd.read_excel(
                        filepath, 
                        sheet_name=sheet_name,
                        nrows=self.max_rows_per_file
                    )
                    
                    # Apply column limit
                    if len(df.columns) > self.max_columns_per_file:
                        logger.warning("excel_column_limit", sheet=sheet_name, columns=len(df.columns), max_columns=self.max_columns_per_file)
                        df = df.iloc[:, :self.max_columns_per_file]
                    
                    # Use first non-empty sheet as primary dataframe
                    if not df.empty and primary_df.empty:
                        primary_df = df
                    
                    sheets[sheet_name] = {
                        'shape': df.shape,
                        'columns': df.columns.tolist(),
                        'data': df.to_dict(orient='records'),
                        'summary': df.describe(include='all').to_dict() if not df.empty else {}
                    }
                    all_dataframes[sheet_name] = df
                    
                except Exception as e:
                    logger.warning("excel_sheet_parse_error", sheet=sheet_name, error=str(e))
                    sheets[sheet_name] = {'error': str(e)}
                    all_dataframes[sheet_name] = pd.DataFrame()
            
            return {
                'type': 'excel',
                'sheets': sheets,
                'sheet_names': xl_file.sheet_names,
                'dataframes': all_dataframes,
                'dataframe': primary_df,  # Primary dataframe for analysis
                'data': primary_df.to_dict(orient='records') if not primary_df.empty else []
            }
        except Exception as e:
            logger.error("excel_parse_error", error=str(e))
            return {'error': str(e), 'type': 'excel'}
    
    async def _parse_html_advanced(self, filepath: Path) -> Dict[str, Any]:
        """Parse HTML with table extraction and row limits"""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Limit HTML content size for very large files
            max_html_size = 10 * 1024 * 1024  # 10MB
            if len(content) > max_html_size:
                logger.warning("html_size_limit", size=len(content), max_size=max_html_size)
                content = content[:max_html_size]
            
            soup = BeautifulSoup(content, 'lxml')
            
            # Extract tables with row limits
            tables = []
            all_tables_data = []
            
            for table in soup.find_all('table'):
                try:
                    df_list = pd.read_html(str(table))
                    if df_list:
                        df = df_list[0]
                        
                        # Apply row limits to HTML tables
                        if len(df) > self.max_rows_per_file:
                            logger.warning("html_table_row_limit", table_rows=len(df), max_rows=self.max_rows_per_file)
                            df = df.head(self.max_rows_per_file)
                        
                        table_data = {
                            'data': df.to_dict(orient='records'),
                            'dataframe': df
                        }
                        tables.append(table_data)
                        all_tables_data.append(df)
                except Exception as e:
                    logger.warning(f"HTML table parse failed: {e}")
                    continue
            
            # Combine all tables into primary dataframe
            primary_df = pd.DataFrame()
            if all_tables_data:
                try:
                    primary_df = pd.concat(all_tables_data, ignore_index=True)
                    # Apply overall row limit
                    if len(primary_df) > self.max_rows_per_file:
                        logger.warning("html_combined_table_limit", combined_rows=len(primary_df), max_rows=self.max_rows_per_file)
                        primary_df = primary_df.head(self.max_rows_per_file)
                except Exception as e:
                    logger.warning("html_table_combine_failed", error=str(e))
                    if all_tables_data:
                        primary_df = all_tables_data[0]
            
            return {
                'type': 'html',
                'text': soup.get_text(separator='\n', strip=True),
                'title': soup.title.string if soup.title else None,
                'tables': tables,
                'table_count': len(tables),
                'links': [{'href': a.get('href'), 'text': a.get_text(strip=True)} 
                         for a in soup.find_all('a', href=True)],
                'dataframe': primary_df,  # Primary dataframe for analysis
                'data': primary_df.to_dict(orient='records') if not primary_df.empty else []
            }
        except Exception as e:
            logger.error("html_parse_error", error=str(e))
            return {'error': str(e), 'type': 'html'}
    
    async def _parse_docx(self, filepath: Path) -> Dict[str, Any]:
        """Parse Word document with table extraction"""
        try:
            doc = docx.Document(filepath)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables with row limits
            tables = []
            all_tables_data = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    if len(table_data) >= self.max_rows_per_file:
                        break
                    table_data.append([cell.text for cell in row.cells])
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    tables.append({
                        'data': df.to_dict(orient='records'), 
                        'dataframe': df
                    })
                    all_tables_data.append(df)
            
            # Combine tables into primary dataframe
            primary_df = pd.DataFrame()
            if all_tables_data:
                try:
                    primary_df = pd.concat(all_tables_data, ignore_index=True)
                    if len(primary_df) > self.max_rows_per_file:
                        primary_df = primary_df.head(self.max_rows_per_file)
                except Exception as e:
                    logger.warning("docx_table_combine_failed", error=str(e))
                    if all_tables_data:
                        primary_df = all_tables_data[0]
            
            return {
                'type': 'docx',
                'paragraphs': paragraphs,
                'full_text': '\n'.join(paragraphs),
                'tables': tables,
                'table_count': len(tables),
                'dataframe': primary_df,
                'data': primary_df.to_dict(orient='records') if not primary_df.empty else []
            }
        except Exception as e:
            logger.error("docx_parse_error", error=str(e))
            return {'error': str(e), 'type': 'docx'}
    
    async def _parse_image_ocr(self, filepath: Path) -> Dict[str, Any]:
        """
        Parse image using OCR (Tesseract/EasyOCR).
        Critical for scanned documents[web:189][web:192].
        """
        try:
            image = Image.open(filepath)
            
            # Perform OCR using Tesseract
            text = pytesseract.image_to_string(image)
            
            logger.info(f"OCR extracted {len(text)} characters from image")
            
            # For OCR results, we don't have tabular data but still return consistent format
            return {
                'type': 'image_ocr',
                'text': text,
                'image_size': image.size,
                'format': image.format,
                'dataframe': pd.DataFrame(),  # Empty dataframe for consistency
                'data': []  # Empty data list for consistency
            }
        except Exception as e:
            logger.error("image_ocr_error", error=str(e))
            return {'error': str(e), 'type': 'image'}
    
    async def _parse_json(self, filepath: Path) -> Dict[str, Any]:
        """Parse JSON file with size limits"""
        try:
            file_size = filepath.stat().st_size
            max_json_size = 10 * 1024 * 1024  # 10MB
            
            if file_size > max_json_size:
                logger.warning("json_size_limit", size=file_size, max_size=max_json_size)
                return {
                    'error': f'JSON file too large ({file_size} bytes), maximum is {max_json_size} bytes',
                    'type': 'json',
                    'dataframe': pd.DataFrame(),
                    'data': []
                }
            
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Try to convert JSON to dataframe if it's tabular
            dataframe = pd.DataFrame()
            if isinstance(data, list) and data and isinstance(data[0], dict):
                try:
                    dataframe = pd.DataFrame(data)
                    # Apply row limits to JSON arrays
                    if len(dataframe) > self.max_rows_per_file:
                        logger.warning("json_row_limit", rows=len(dataframe), max_rows=self.max_rows_per_file)
                        dataframe = dataframe.head(self.max_rows_per_file)
                        data = dataframe.to_dict(orient='records')
                except Exception as e:
                    logger.warning("json_to_dataframe_failed", error=str(e))
            
            return {
                'type': 'json', 
                'data': data,
                'dataframe': dataframe
            }
        except Exception as e:
            logger.error("json_parse_error", error=str(e))
            return {'error': str(e), 'type': 'json'}
    
    async def _parse_jsonl(self, filepath: Path) -> Dict[str, Any]:
        """Parse JSON Lines file with row limits"""
        try:
            data = []
            line_count = 0
            
            with open(filepath, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        if line_count >= self.max_rows_per_file:
                            logger.warning("jsonl_row_limit", rows=line_count, max_rows=self.max_rows_per_file)
                            break
                        data.append(json.loads(line))
                        line_count += 1
            
            dataframe = pd.DataFrame()
            if data and isinstance(data[0], dict):
                try:
                    dataframe = pd.DataFrame(data)
                except Exception as e:
                    logger.warning("jsonl_to_dataframe_failed", error=str(e))
            
            return {
                'type': 'jsonl', 
                'data': data, 
                'count': len(data),
                'dataframe': dataframe
            }
        except Exception as e:
            logger.error("jsonl_parse_error", error=str(e))
            return {'error': str(e), 'type': 'jsonl'}
    
    async def _parse_xml(self, filepath: Path) -> Dict[str, Any]:
        """Parse XML file"""
        try:
            import xml.etree.ElementTree as ET
            tree = ET.parse(filepath)
            root = tree.getroot()
            
            def xml_to_dict(element):
                result = {}
                for child in element:
                    if len(child) > 0:
                        result[child.tag] = xml_to_dict(child)
                    else:
                        result[child.tag] = child.text
                return result
            
            data = xml_to_dict(root)
            
            return {
                'type': 'xml',
                'data': data,
                'root_tag': root.tag,
                'dataframe': pd.DataFrame(),  # XML typically not tabular
                'data': [data]  # Wrap in list for consistency
            }
        except Exception as e:
            logger.error("xml_parse_error", error=str(e))
            return {'error': str(e), 'type': 'xml'}
    
    async def _parse_yaml(self, filepath: Path) -> Dict[str, Any]:
        """Parse YAML file"""
        try:
            import yaml
            with open(filepath, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
            
            dataframe = pd.DataFrame()
            if isinstance(data, list) and data and isinstance(data[0], dict):
                try:
                    dataframe = pd.DataFrame(data)
                    if len(dataframe) > self.max_rows_per_file:
                        dataframe = dataframe.head(self.max_rows_per_file)
                        data = dataframe.to_dict(orient='records')
                except Exception as e:
                    logger.warning("yaml_to_dataframe_failed", error=str(e))
            
            return {
                'type': 'yaml', 
                'data': data,
                'dataframe': dataframe
            }
        except Exception as e:
            logger.error("yaml_parse_error", error=str(e))
            return {'error': str(e), 'type': 'yaml'}
    
    async def _parse_text(self, filepath: Path) -> Dict[str, Any]:
        """Parse plain text file with size limits"""
        try:
            file_size = filepath.stat().st_size
            max_text_size = 5 * 1024 * 1024  # 5MB
            
            if file_size > max_text_size:
                logger.warning("text_size_limit", size=file_size, max_size=max_text_size)
                # Read only first 5MB
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read(max_text_size)
                text += f"\n\n[FILE TRUNCATED - ORIGINAL SIZE: {file_size} bytes]"
            else:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            
            return {
                'type': 'text', 
                'text': text,
                'dataframe': pd.DataFrame(),  # Text files don't have tabular data
                'data': []  # Empty data list for consistency
            }
        except Exception as e:
            logger.error("text_parse_error", error=str(e))
            return {'error': str(e), 'type': 'text'}