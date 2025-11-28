# =============================================================================
# FILE: src/parsers/docx_parser.py
# =============================================================================
"""
Advanced DOCX (Microsoft Word) parser with support for paragraphs, tables,
images, headers, footers, and metadata extraction.
References: [web:264][web:265][web:268][web:271]
"""

import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from pathlib import Path
from typing import Dict, Any, List, Optional
import structlog
import pandas as pd
import base64
from io import BytesIO

logger = structlog.get_logger()


class DOCXParser:
    """
    Parse Microsoft Word (.docx) documents with comprehensive extraction
    of text, tables, images, styles, and metadata.
    """
    
    @staticmethod
    def parse_file(filepath: Path) -> Dict[str, Any]:
        """
        Parse DOCX file and extract all content[web:265][web:268][web:271].
        
        Args:
            filepath: Path to DOCX file
            
        Returns:
            Dict containing all extracted content and metadata
        """
        try:
            doc = docx.Document(str(filepath))
            
            # Extract various components
            paragraphs = DOCXParser._extract_paragraphs(doc)
            tables = DOCXParser._extract_tables(doc)
            images = DOCXParser._extract_images(doc)
            metadata = DOCXParser._extract_metadata(doc)
            headers_footers = DOCXParser._extract_headers_footers(doc)
            
            # Get full text
            full_text = '\n'.join([p['text'] for p in paragraphs if p['text']])
            
            logger.info(
                "docx_parsed",
                filepath=str(filepath),
                paragraphs=len(paragraphs),
                tables=len(tables),
                images=len(images)
            )
            
            return {
                'type': 'docx',
                'filepath': str(filepath),
                'full_text': full_text,
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'tables': tables,
                'table_count': len(tables),
                'images': images,
                'image_count': len(images),
                'headers_footers': headers_footers,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error("docx_parse_error", filepath=str(filepath), error=str(e))
            return {
                'error': str(e),
                'type': 'docx',
                'filepath': str(filepath)
            }
    
    @staticmethod
    def _extract_paragraphs(doc: docx.Document) -> List[Dict[str, Any]]:
        """
        Extract paragraphs with text and style information[web:268][web:271].
        
        Args:
            doc: Document object
            
        Returns:
            List of paragraph dicts
        """
        paragraphs = []
        
        for para in doc.paragraphs:
            paragraphs.append({
                'text': para.text,
                'style': para.style.name if para.style else None,
                'alignment': str(para.alignment) if para.alignment else None,
                'is_heading': para.style.name.startswith('Heading') if para.style else False
            })
        
        return paragraphs
    
    @staticmethod
    def _extract_tables(doc: docx.Document) -> List[Dict[str, Any]]:
        """
        Extract tables as both structured data and DataFrames[web:268].
        
        Args:
            doc: Document object
            
        Returns:
            List of table dicts with data and DataFrames
        """
        tables = []
        
        for table in doc.tables:
            # Extract table data
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text (including nested paragraphs)
                    cell_text = '\n'.join([p.text for p in cell.paragraphs])
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            # Convert to DataFrame if possible
            df = None
            if table_data and len(table_data) > 1:
                try:
                    # Use first row as headers
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                except Exception as e:
                    logger.debug(f"Could not convert table to DataFrame: {e}")
                    df = pd.DataFrame(table_data)
            
            tables.append({
                'data': table_data,
                'dataframe': df,
                'rows': len(table.rows),
                'columns': len(table.columns)
            })
        
        return tables
    
    @staticmethod
    def _extract_images(doc: docx.Document) -> List[Dict[str, Any]]:
        """
        Extract embedded images from document[web:265].
        
        Args:
            doc: Document object
            
        Returns:
            List of image dicts with base64-encoded data
        """
        images = []
        
        # Access document relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # Encode as base64 for transport
                    image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    
                    images.append({
                        'filename': rel.target_ref.split('/')[-1],
                        'content_type': image_part.content_type,
                        'size_bytes': len(image_bytes),
                        'data_base64': image_base64
                    })
                except Exception as e:
                    logger.warning(f"Failed to extract image: {e}")
        
        return images
    
    @staticmethod
    def _extract_metadata(doc: docx.Document) -> Dict[str, Any]:
        """
        Extract document metadata (author, created date, etc.)[web:265].
        
        Args:
            doc: Document object
            
        Returns:
            Dict containing metadata
        """
        try:
            core_props = doc.core_properties
            
            return {
                'author': core_props.author,
                'title': core_props.title,
                'subject': core_props.subject,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
                'keywords': core_props.keywords,
                'comments': core_props.comments
            }
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {e}")
            return {}
    
    @staticmethod
    def _extract_headers_footers(doc: docx.Document) -> Dict[str, List[str]]:
        """
        Extract text from headers and footers[web:265].
        
        Args:
            doc: Document object
            
        Returns:
            Dict with header and footer texts
        """
        headers = []
        footers = []
        
        for section in doc.sections:
            # Headers
            for para in section.header.paragraphs:
                if para.text.strip():
                    headers.append(para.text)
            
            # Footers
            for para in section.footer.paragraphs:
                if para.text.strip():
                    footers.append(para.text)
        
        return {
            'headers': headers,
            'footers': footers
        }
    @staticmethod
    def _generate_summary(df: pd.DataFrame) -> Dict[str, Any]:
        """
        Generate summary statistics for a sheet"""
        summary = {}
        
        # Numeric summary
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            summary['numeric'] = df[numeric_cols].describe().to_dict()
        
        # Categorical summary
        categorical_cols = df.select_dtypes(include=['object']).columns
        if len(categorical_cols) > 0:
            summary['categorical'] = {col: df[col].value_counts().to_dict() for col in categorical_cols}

        return summary
    